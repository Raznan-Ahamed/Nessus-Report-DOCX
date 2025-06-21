import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK
from collections import defaultdict
import matplotlib.pyplot as plt
import os
import logging

logging.basicConfig(
    level=logging.INFO,
    format='[%(levelname)s] %(message)s'
)

def get_severity_color(severity):
    severity = severity.upper()
    if severity in ("CRITICAL", "HIGH"):
        return "FF0000"
    elif severity == "MEDIUM":
        return "FFA500"
    elif severity == "LOW":
        return "FFFF00"
    return "D3D3D3"

def set_cell_background(cell, rgb_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)

def add_heading(document, text, level):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 13 if level == 2 else 12)
    run.font.name = 'Aptos'
    return p

def add_paragraph(cell, text, bold=False, size=12, color=None):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Aptos'
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'nil')
        borders.append(elem)
    tblPr.append(borders)

def add_severity_chart(df, document, chart_name="severity_chart.png", title="Vulnerabilities by Severity"):
    severity_counts = df['Risk'].value_counts().reindex(['CRITICAL', 'MEDIUM', 'LOW'], fill_value=0)

    plt.figure(figsize=(5, 3))
    bars = plt.bar(severity_counts.index, severity_counts.values, color=['red', 'orange', 'yellow'])
    plt.title(title)
    plt.xlabel('Severity')
    plt.ylabel('Count')
    plt.tight_layout()

    plt.savefig(chart_name)
    plt.close()

    document.add_picture(chart_name)
    os.remove(chart_name)

def generate_report(csv_path, template_path, output_path):
    df = pd.read_csv(csv_path)
    logging.info(f"Loaded CSV: {csv_path}")
    logging.info(f"Using DOCX template: {template_path}")
    df = df[df['Risk'].notna()]
    df = df[~df['Risk'].str.lower().eq('none')]
    df['NormalizedRisk'] = df['Risk'].str.upper().replace({'HIGH': 'CRITICAL'})
    df = df[['Host', 'NormalizedRisk', 'Name', 'Description', 'Solution']]
    df.columns = ['Host', 'Risk', 'Name', 'Description', 'Solution']

    grouped = defaultdict(lambda: defaultdict(list))
    for _, row in df.iterrows():
        grouped[row['Host']][row['Risk']].append({
            'title': row['Name'],
            'description': row['Description'],
            'solution': row['Solution'],
            'risk': row['Risk']
        })
    logging.info(f"Total vulnerabilities after filtering: {len(df)}")

    doc = Document(template_path)
    doc.add_page_break()  # page 3 start

    doc.add_paragraph("1. Statistics", style='Heading 1')
    add_severity_chart(df, doc)
    doc.add_paragraph("\n", style='Normal')
    doc.add_page_break()

    doc.add_paragraph("2. Executive Summary", style='Heading 1')
    doc.add_paragraph("Add executive summary here...\n", style='Normal')
    doc.add_page_break()

    doc.add_paragraph("3. Vulnerabilities and Remediations", style='Heading 1')

    severity_order = ['CRITICAL', 'MEDIUM', 'LOW']

    host_count = 0
    for host, severities in grouped.items():
        host_count += 1
        logging.info(f"Processing host: {host}")
        doc.add_paragraph(f"3.{host_count} {host}", style='Heading 2')
            # Filter the DataFrame for the current host
        host_df = df[df['Host'] == host]
        chart_file = f"{host.replace('.', '_')}_chart.png"
        add_severity_chart(host_df, doc, chart_name=chart_file, title=f"{host} - Vulnerabilities by Severity")


        for i, severity in enumerate(severity_order, start=1):
            if severity not in severities:
                continue

            vulns = severities[severity]
            logging.info(f"  Adding {severity} vulnerabilities: {len(vulns)} entries")
            doc.add_paragraph(f"3.{host_count}.{i} {severity.capitalize()} Vulnerabilities", style='Heading 3')

            for j, vuln in enumerate(vulns, start=1):
                doc.add_paragraph(f"3.{host_count}.{i}.{j} {vuln['title']}", style='Heading 4')
                #doc.add_page_break()

                table = doc.add_table(rows=5, cols=1)
                table.autofit = True
                remove_table_borders(table)

                # Title with background color
                cell = table.cell(0, 0)
                set_cell_background(cell, get_severity_color(vuln['risk']))
                add_paragraph(cell, vuln['title'], bold=True, size=16, color="FFFFFF")

                # Risk
                add_paragraph(table.cell(1, 0), f"Risk: {vuln['risk']}", bold=True)

                # Description
                desc_cell = table.cell(2, 0)
                add_paragraph(desc_cell, "Description:", bold=True)
                desc_cell.add_paragraph(vuln['description'])

                # Impact
                impact_cell = table.cell(3, 0)
                add_paragraph(impact_cell, "Impact:", bold=True)
                impact_cell.add_paragraph("Impact information not available.")

                # Solution
                sol_cell = table.cell(4, 0)
                set_cell_background(sol_cell, "9ACD32")
                add_paragraph(sol_cell, "Remediation:", bold=True)
                sol_cell.add_paragraph(vuln['solution'])

                doc.add_page_break()

    doc.save(output_path)
    logging.info(f"Report saved to: {output_path}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Generate a structured vulnerability report from Nessus CSV.")
    parser.add_argument("csv_file", help="Path to the Nessus CSV file")
    parser.add_argument("--template", default="REPORT_TEMPLATE.docx", help="Path to the DOCX template")
    parser.add_argument("--output", default="nessus_vuln_report.docx", help="Output DOCX file name")
    args = parser.parse_args()

    generate_report(args.csv_file, args.template, args.output)
